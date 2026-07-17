# -*- coding: utf-8 -*-
"""
一键处理：将新docx文档添加到两个表格中
  - 表格1（小说正文1_617.xlsx）：标题 + 正文（含提醒语+标签）
  - 表格2（CLI_六层梗_617.xlsx）：洗后梗 + 标题 + 纯正文 + 选项
"""

import os
import re
import json
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

# ============================================================
# 新文档路径
# ============================================================
NEW_DOCX = r"D:/小说/生文合集3/易遇只是变成鬼了，又不是不爱姑姑了，舍不得看到姑姑难过的.docx"

TABLE1_XLSX = r"D:/小说/表格/小说正文1_617.xlsx"
TABLE2_XLSX = r"D:/小说/CLI_六层梗_617.xlsx"

# ============================================================
# 国乙男主表
# ============================================================
GAME_CHARACTER_MAP = {
    "恋与制作人": ["李泽言", "许墨", "白起", "周棋洛"],
    "光与夜之恋": ["陆沉", "齐司礼", "萧逸", "查理苏", "夏鸣星"],
    "未定事件簿": ["左然", "莫弈", "夏彦", "陆景和"],
    "时空中的绘旅人": ["叶瑄", "司岚", "路辰", "罗夏", "艾因"],
    "恋与深空": ["秦彻", "祁煜", "沈星回", "黎深", "夏以昼"],
    "世界之外": ["柏源", "夏萧因", "易遇", "顾时夜"],
}
CHAR_TO_GAME = {}
for game, chars in GAME_CHARACTER_MAP.items():
    for c in chars:
        CHAR_TO_GAME[c] = game
SORTED_CHARS = sorted(CHAR_TO_GAME.keys(), key=len, reverse=True)

REMINDER_TEXT = "提醒一下宝宝们不用蹲，后续4连进群⬇️自取就可以啦~不会在这里更新哦 ᗜ ‸ ᗜ"
SEPARATOR_PATTERN = re.compile(r'^-{3,}$')
TAG_PATTERN = re.compile(r'^#[^\s]+(\s+#[^\s]+)+$')


def detect_character_and_game(title):
    for char in SORTED_CHARS:
        if char in title:
            return char, CHAR_TO_GAME[char]
    return None, None


def read_docx(filepath):
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paruments if hasattr(doc, 'paruments') else doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return '\n'.join(paragraphs)


def clean_content(raw_content):
    lines = raw_content.split('\n')
    remove_indices = set()
    i = len(lines) - 1
    while i >= 0 and lines[i].strip() == '':
        remove_indices.add(i); i -= 1
    if i >= 0 and TAG_PATTERN.match(lines[i].strip()):
        remove_indices.add(i); i -= 1
    while i >= 0 and lines[i].strip() == '':
        remove_indices.add(i); i -= 1
    if i >= 0 and REMINDER_TEXT in lines[i]:
        remove_indices.add(i); i -= 1
    while i >= 0 and lines[i].strip() == '':
        remove_indices.add(i); i -= 1
    if i >= 0 and SEPARATOR_PATTERN.match(lines[i].strip()):
        remove_indices.add(i); i -= 1
    while i >= 0 and lines[i].strip() == '':
        remove_indices.add(i); i -= 1
    cleaned = [lines[j] for j in range(len(lines)) if j not in remove_indices]
    return '\n'.join(cleaned).rstrip('\n')


# ============================================================
# 新小说的六层梗分析和选项
# ============================================================
NEW_NOVEL_ANALYSIS = {
    "six_layers": "第一层·噱头层：灯语传情+追悼会醋意\n第二层·世界观层：现代都市/灵异\n第三层·人设层：鬼魂养侄/寡姑姑/醋精亡灵\n第四层·情节层：死亡+灯语传信+追悼会守护+备忘录对话\n第五层·情绪层：丧亲虐+守护甜+醋意爽\n第六层·结局层：阴阳相守",
    "options": "A. 你对着备忘录打字：\"你说你看着我，那你现在过来。别用灯，用别的。\"\nB. 你把手机翻过去扣在桌上，声音很轻：\"小遇，你别守着我了。你该走了。\"\nC. 你在备忘录里打了一行字删掉，又打了一行又删掉，最后只留了三个字：\"想你了。\""
}


def process_new_docx():
    """处理新docx，添加到两个表格"""

    # ---- 读取docx ----
    doc = Document(NEW_DOCX)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    raw_content = '\n'.join(paragraphs)
    title = os.path.splitext(os.path.basename(NEW_DOCX))[0]

    # 去掉正文开头可能重复的标题行
    if raw_content.startswith(title):
        raw_content = raw_content[len(title):].lstrip('\n')

    print(f"标题: {title}")
    print(f"正文字数: {len(raw_content)}")

    # 匹配角色
    char_name, game_name = detect_character_and_game(title)
    print(f"角色匹配: {char_name} -> {game_name}")

    # ---- 表格1：添加带提醒语+标签的正文 ----
    print("\n=== 处理表格1 ===")
    wb1 = load_workbook(TABLE1_XLSX)
    ws1 = wb1.active

    # 找到下一个空行
    next_row = 2
    while ws1.cell(row=next_row, column=1).value:
        next_row += 1

    # 构建带提醒语+标签的正文
    content_with_tags = raw_content + "\n\n" + "-" * 70 + "\n" + REMINDER_TEXT + "\n\n"
    if char_name and game_name:
        content_with_tags += "#" + char_name + " #" + game_name + " #bg #乙女向 #同人文"
    else:
        content_with_tags += "#bg #乙女向 #同人文"

    ws1.cell(row=next_row, column=1, value=title)
    ws1.cell(row=next_row, column=2, value=content_with_tags)
    print(f"  添加到第 {next_row} 行")

    wb1.save(TABLE1_XLSX)
    print(f"  表格1已保存: {TABLE1_XLSX}")

    # ---- 表格2：添加洗后梗+纯正文+选项 ----
    print("\n=== 处理表格2 ===")
    wb2 = load_workbook(TABLE2_XLSX)
    ws2 = wb2.active

    next_row2 = 2
    while ws2.cell(row=next_row2, column=2).value:
        next_row2 += 1

    # 清洗正文
    clean_text = clean_content(raw_content)

    # 洗后梗
    if game_name:
        world_prefix = "世界观 " + game_name + "同人"
    else:
        world_prefix = "世界观 未知同人"
    a_content = world_prefix + "\n" + NEW_NOVEL_ANALYSIS["six_layers"]

    # 样式
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin'),
    )
    content_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    cell_a = ws2.cell(row=next_row2, column=1, value=a_content)
    cell_a.alignment = content_alignment
    cell_a.border = thin_border
    cell_a.font = Font(size=10)

    cell_b = ws2.cell(row=next_row2, column=2, value=title)
    cell_b.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_b.border = thin_border
    cell_b.font = Font(size=10)

    cell_c = ws2.cell(row=next_row2, column=3, value=clean_text)
    cell_c.alignment = content_alignment
    cell_c.border = thin_border
    cell_c.font = Font(size=10)

    cell_d = ws2.cell(row=next_row2, column=4, value=NEW_NOVEL_ANALYSIS["options"])
    cell_d.alignment = content_alignment
    cell_d.border = thin_border
    cell_d.font = Font(size=10)

    print(f"  添加到第 {next_row2} 行")

    wb2.save(TABLE2_XLSX)
    print(f"  表格2已保存: {TABLE2_XLSX}")

    print(f"\n{'='*50}")
    print(f"新文档处理完成！")
    print(f"  表格1: {TABLE1_XLSX} (第{next_row}行)")
    print(f"  表格2: {TABLE2_XLSX} (第{next_row2}行)")


if __name__ == "__main__":
    process_new_docx()
